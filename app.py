import streamlit as st
import matplotlib.pyplot as plt
import pandas as pd
import io
import re
import nltk
import spacy
from collections import Counter
import PyPDF2
import docx
from rank_bm25 import BM25Okapi
import traceback
import logging

from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document

from sentence_transformers import SentenceTransformer
from langchain_community.embeddings import HuggingFaceEmbeddings

# Setup logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

# ---------------------------------------
# Configuration and Constants
SKILL_KEYWORDS = [
    "python", "r", "sql", "excel", "tableau", "power bi", "powerbi",
    "machine learning", "data science", "deep learning", "nlp", "ai",
    "java", "javascript", "c++", "c#", "html", "css", "react", "nodejs",
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "matplotlib",
    "statistics", "analytics", "visualization", "hadoop", "spark", "kafka",
    "docker", "kubernetes", "aws", "azure", "gcp", "mongodb", "postgresql",
    "git", "github", "agile", "scrum", "project management"
]

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MIN_CONTENT_LENGTH = 50
MIN_JD_LENGTH = 20

# ---------------------------------------
# Initialization and dependency checks
@st.cache_resource
def load_dependencies():
    """Load and cache dependencies"""
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        st.error("❌ spaCy model 'en_core_web_sm' not found. Install with:\n\n`python -m spacy download en_core_web_sm`")
        st.stop()
    
    try:
        nltk.download('punkt', quiet=True)
    except Exception as e:
        st.error(f"❌ Failed to download NLTK data: {str(e)}")
        st.stop()
    
    return nlp

nlp = load_dependencies()

# ---------------------------------------
# Helper Functions
def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    # Remove special characters but keep spaces
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.lower().strip()

def extract_skills(text):
    """Extract skills from text using NLP and keyword matching"""
    if not text:
        return []
    
    try:
        doc = nlp(text)
        # Extract entities that could be skills (organizations, products)
        skills_found = []
        for ent in doc.ents:
            if ent.label_ in ["ORG", "PRODUCT", "GPE"] and len(ent.text) > 2:
                skills_found.append(ent.text.lower().strip())
        
        # Add keyword matching for technical skills
        text_lower = text.lower()
        for skill in SKILL_KEYWORDS:
            if re.search(rf"\b{re.escape(skill)}\b", text_lower):
                skills_found.append(skill)
        
        # Remove duplicates and filter out common words
        skills_found = list(set(skills_found))
        # Filter out very short or common words
        skills_found = [skill for skill in skills_found if len(skill) > 2 and skill not in ['the', 'and', 'for', 'with', 'data']]
        
        return skills_found
    except Exception as e:
        logger.warning(f"Error extracting skills: {str(e)}")
        return []

def extract_experience(text):
    """Extract years of experience from text"""
    if not text:
        return 0
    
    try:
        patterns = [
            r'(\d+)\s*\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\s*-\s*(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'experience:?\s*(\d+)\s*(?:years?|yrs?)',
            r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\s*\+\s*(?:years?|yrs?)',
            r'over\s*(\d+)\s*(?:years?|yrs?)',
            r'more\s*than\s*(\d+)\s*(?:years?|yrs?)'
        ]
        
        years = []
        text_lower = text.lower()
        
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):
                    # For range patterns (e.g., "3-5 years"), take the maximum
                    valid_years = [int(x) for x in match if x.isdigit() and 0 < int(x) < 50]
                    if valid_years:
                        years.extend(valid_years)
                else:
                    if match.isdigit() and 0 < int(match) < 50:
                        years.append(int(match))
        
        return max(years) if years else 0
    except Exception as e:
        logger.warning(f"Error extracting experience: {str(e)}")
        return 0

def parse_sections(text):
    """Parse text into sections (skills, experience, etc.)"""
    if not text:
        return {"skills": [], "experience": [], "education": []}
    
    sections = {"skills": [], "experience": [], "education": []}
    lines = text.split('\n')
    
    for line in lines:
        line_lower = line.lower().strip()
        if not line_lower:
            continue
            
        # Check for skills sections
        if any(keyword in line_lower for keyword in ['skill', 'technical', 'competenc', 'proficien']):
            sections["skills"].append(line)
        
        # Check for experience sections
        elif any(keyword in line_lower for keyword in ['experience', 'work', 'employ', 'career', 'position']):
            sections["experience"].append(line)
        
        # Check for education sections
        elif any(keyword in line_lower for keyword in ['education', 'degree', 'university', 'college', 'graduat']):
            sections["education"].append(line)
    
    # Add fallback if no experience sections found
    if not sections["experience"]:
        sections["experience"].append(text[:200])  # Use first 200 chars as fallback
    
    return sections

def extract_pdf_text(file):
    """Extract text from PDF file"""
    try:
        text = ""
        reader = PyPDF2.PdfReader(file)
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Error reading page {page_num + 1} of {file.name}: {str(e)}")
                continue
        
        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF {file.name}: {str(e)}")
        return ""

def extract_docx_text(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file)
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        return "\n".join(paragraphs)
    except Exception as e:
        st.error(f"Error reading DOCX {file.name}: {str(e)}")
        return ""

def create_weighted_content(original_text, sections):
    """Create weighted content by emphasizing important sections"""
    content_weighted = clean_text(original_text)
    
    # Add extra weight to skills sections
    for skill_section in sections["skills"]:
        clean_skill = clean_text(skill_section)
        if clean_skill:
            content_weighted += (" " + clean_skill) * 2
    
    # Add extra weight to experience sections
    for exp_section in sections["experience"]:
        clean_exp = clean_text(exp_section)
        if clean_exp:
            content_weighted += (" " + clean_exp) * 2
    
    return content_weighted

def validate_file_size(file):
    """Validate file size"""
    try:
        if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
            return False, f"File too large (>{MAX_FILE_SIZE/1024/1024:.1f}MB)"
        return True, ""
    except AttributeError:
        return True, ""  # Size check not available, proceed

def process_resume(file):
    """Process a single resume file"""
    try:
        # Validate file size
        is_valid, error_msg = validate_file_size(file)
        if not is_valid:
            return None, error_msg
        
        # Extract text based on file type
        if file.name.lower().endswith('.pdf'):
            text_content = extract_pdf_text(file)
        elif file.name.lower().endswith('.docx'):
            text_content = extract_docx_text(file)
        else:
            return None, "Unsupported file format"
        
        # Validate content length
        if not text_content or len(text_content.strip()) < MIN_CONTENT_LENGTH:
            return None, "Content too short or empty"
        
        # Extract information
        skills = extract_skills(text_content)
        experience = extract_experience(text_content)
        sections = parse_sections(text_content)
        weighted_content = create_weighted_content(text_content, sections)
        
        resume_data = {
            "name": file.name,
            "original_text": text_content,
            "weighted_content": weighted_content,
            "skills": skills,
            "experience": experience,
            "sections": sections
        }
        
        return resume_data, ""
    except Exception as e:
        error_msg = f"Error processing file: {str(e)}"
        logger.error(error_msg)
        return None, error_msg

def build_bm25_index(resume_data):
    """Build BM25 index from resume data"""
    try:
        tokenized_corpus = []
        for resume in resume_data:
            tokens = nltk.word_tokenize(resume["weighted_content"])
            tokenized_corpus.append(tokens)
        
        bm25 = BM25Okapi(tokenized_corpus)
        return bm25
    except Exception as e:
        st.error(f"Error building BM25 index: {str(e)}")
        return None

def build_faiss_index(resume_data):
    """Build FAISS vector index from resume data"""
    try:
        # Initialize embeddings with correct parameters
        model_name = "sentence-transformers/all-MiniLM-L6-v2"
        embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        documents = []
        for resume in resume_data:
            doc = Document(
                page_content=resume["weighted_content"],
                metadata={"name": resume["name"]}
            )
            documents.append(doc)
        
        vector_store = FAISS.from_documents(documents, embeddings)
        return vector_store
    except Exception as e:
        st.error(f"Error building FAISS index: {str(e)}")
        st.error(f"Traceback: {traceback.format_exc()}")
        return None

def calculate_scores(resume_data, bm25_index, vector_store, job_description):
    """Calculate combined BM25 and semantic similarity scores"""
    try:
        jd_clean = clean_text(job_description)
        jd_tokens = nltk.word_tokenize(jd_clean)
        
        # Get BM25 scores
        bm25_scores = bm25_index.get_scores(jd_tokens)
        
        # Normalize BM25 scores with better handling of edge cases
        if not bm25_scores:
            bm25_normalized = [50.0] * len(resume_data)
        else:
            bm_min, bm_max = min(bm25_scores), max(bm25_scores)
            bm_range = max(bm_max - bm_min, 0.001)  # Prevent division by zero
            
            # Scale to 0-100 range with minimum of 10% to avoid 0% scores
            bm25_normalized = [10 + 90 * ((score - bm_min) / bm_range) for score in bm25_scores]
        
        # Get FAISS similarity scores with better distance conversion
        faiss_results = vector_store.similarity_search_with_score(job_description, k=len(resume_data))
        
        # Create mapping from document name to original index
        name_to_index = {resume["name"]: idx for idx, resume in enumerate(resume_data)}
        
        # Combine scores
        results = []
        for doc, semantic_distance in faiss_results:
            doc_name = doc.metadata["name"]
            
            # Find original index
            original_idx = name_to_index.get(doc_name)
            if original_idx is None:
                continue
            
            # Get original resume data
            original_resume = resume_data[original_idx]
            
            # Get BM25 score for this document
            bm25_score = bm25_normalized[original_idx]
            
            # Convert semantic distance to similarity percentage (0-100)
            # Adjusted conversion to avoid 0% scores
            semantic_score = 100 * (1 - min(semantic_distance, 1.0))  # Cap distance at 1.0
            semantic_score = max(10, semantic_score)  # Ensure minimum 10% score
            
            # Combine scores with adjusted weights
            final_score = round(0.5 * bm25_score + 0.5 * semantic_score, 2)
            
            results.append({
                "name": doc_name,
                "final_score": final_score,
                "bm25_score": round(bm25_score, 2),
                "semantic_score": round(semantic_score, 2),
                "skills": original_resume["skills"],
                "experience": original_resume["experience"]
            })
        
        # Sort by final score (descending)
        results.sort(key=lambda x: x["final_score"], reverse=True)
        return results
    
    except Exception as e:
        st.error(f"Error calculating scores: {str(e)}")
        return []

def calculate_scores(resume_data, bm25_index, vector_store, job_description):
    """Calculate combined BM25 and semantic similarity scores"""
    try:
        jd_clean = clean_text(job_description)
        jd_tokens = nltk.word_tokenize(jd_clean)
        
        # Get BM25 scores
        bm25_scores = bm25_index.get_scores(jd_tokens)
        
        # Normalize BM25 scores
        bm_min, bm_max = min(bm25_scores), max(bm25_scores)
        bm_range = bm_max - bm_min
        
        if bm_range == 0:
            st.warning("⚠️ BM25 scores are uniform, results may be less meaningful.")
            bm25_normalized = [50.0] * len(bm25_scores)
        else:
            bm25_normalized = [((score - bm_min) / bm_range) * 100 for score in bm25_scores]
        
        # Get FAISS similarity scores
        faiss_results = vector_store.similarity_search_with_score(job_description, k=len(resume_data))
        
        # Create mapping from document name to original index
        name_to_index = {resume["name"]: idx for idx, resume in enumerate(resume_data)}
        
        # Combine scores
        results = []
        for doc, semantic_distance in faiss_results:
            doc_name = doc.metadata["name"]
            
            # Find original index
            original_idx = name_to_index.get(doc_name)
            if original_idx is None:
                st.error(f"Document mapping error for {doc_name}")
                continue
            
            # Get original resume data
            original_resume = resume_data[original_idx]
            
            # Get BM25 score for this document
            bm25_score = bm25_normalized[original_idx]
            
            # Convert semantic distance to similarity percentage
            semantic_score = 100 - (semantic_distance * 20)
            semantic_score = max(0, min(100, semantic_score))
            
            # Combine scores (you can adjust weights here)
            final_score = round(0.6 * bm25_score + 0.4 * semantic_score, 2)
            
            results.append({
                "name": doc_name,
                "final_score": final_score,
                "bm25_score": round(bm25_score, 2),
                "semantic_score": round(semantic_score, 2),
                "skills": original_resume["skills"],
                "experience": original_resume["experience"]
            })
        
        # Sort by final score (descending)
        results.sort(key=lambda x: x["final_score"], reverse=True)
        return results
    
    except Exception as e:
        st.error(f"Error calculating scores: {str(e)}")
        st.error(f"Traceback: {traceback.format_exc()}")
        return []

def display_results(results):
    """Display results in Streamlit"""
    if not results:
        st.error("No results to display")
        return
    
    st.subheader("✅ Resume Matching Results")
    
    # Display detailed results
    for i, result in enumerate(results):
        with st.expander(f"**{i+1}. {result['name']}** - {result['final_score']}% match"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("Final Score", f"{result['final_score']}%")
                st.metric("BM25 Score", f"{result['bm25_score']}%")
                st.metric("Semantic Score", f"{result['semantic_score']}%")
            
            with col2:
                st.metric("Experience", f"{result['experience']} years")
                if result['skills']:
                    st.write("**Skills:**", ", ".join(result['skills'][:10]))  # Show first 10 skills
                else:
                    st.write("**Skills:** None detected")

def create_visualization(results):
    """Create visualization of results"""
    if not results:
        return
    
    names = [result['name'] for result in results]
    scores = [result['final_score'] for result in results]
    
    # Create horizontal bar chart
    fig, ax = plt.subplots(figsize=(12, max(6, len(names) * 0.6)))
    
    # Create color gradient based on scores
    colors = plt.cm.RdYlGn([score/100 for score in scores])
    
    bars = ax.barh(names, scores, color=colors)
    ax.set_xlabel("Match Score (%)")
    ax.set_title("Resume Matching Results", fontsize=16, fontweight='bold')
    ax.set_xlim(0, 100)
    
    # Add score labels on bars
    for bar, score in zip(bars, scores):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2, 
                f'{score}%', ha='left', va='center', fontweight='bold')
    
    ax.invert_yaxis()
    plt.tight_layout()
    st.pyplot(fig)

def create_excel_download(results):
    """Create Excel download"""
    if not results:
        return
    
    # Prepare data for Excel
    excel_data = []
    for i, result in enumerate(results):
        excel_data.append({
            "Rank": i + 1,
            "Resume Name": result['name'],
            "Final Score (%)": result['final_score'],
            "BM25 Score (%)": result['bm25_score'],
            "Semantic Score (%)": result['semantic_score'],
            "Skills": ', '.join(result['skills']) if result['skills'] else 'None',
            "Experience (Years)": result['experience']
        })
    
    df = pd.DataFrame(excel_data)
    
    # Create Excel buffer
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name="Resume_Analysis")
        
        # Get workbook and worksheet
        workbook = writer.book
        worksheet = writer.sheets['Resume_Analysis']
        
        # Auto-adjust column widths
        for column_cells in worksheet.iter_cols():
            max_length = 0
            column_letter = column_cells[0].column_letter
            for cell in column_cells:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    st.download_button(
        label="📥 Download Detailed Results (Excel)",
        data=buffer.getvalue(),
        file_name="resume_analysis_results.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# ---------------------------------------
# Main Streamlit Application
def main():
    st.set_page_config(page_title="Resume Matcher", page_icon="📄", layout="wide")
    
    st.title("📄 AI-Powered Resume Matcher")
    st.markdown("### Match resumes against job descriptions using hybrid BM25 + Semantic similarity")
    
    # Input section
    col1, col2 = st.columns([2, 1])
    
    with col1:
        job_description = st.text_area(
            "Job Description",
            height=200,
            placeholder="Paste the complete job description here...",
            help="Enter a detailed job description including required skills, experience, and qualifications"
        )
    
    with col2:
        st.markdown("### Upload Resumes")
        uploaded_files = st.file_uploader(
            "Select PDF or DOCX files",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            help="Upload multiple resume files for comparison"
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files uploaded")
    
    # Validation
    if not job_description.strip():
        st.info("👆 Please enter a job description to start matching")
        return
    
    if len(job_description.strip()) < MIN_JD_LENGTH:
        st.warning(f"Job description should be at least {MIN_JD_LENGTH} characters long")
        return
    
    if not uploaded_files:
        st.info("👆 Please upload resume files to analyze")
        return
    
    # Process button
    if st.button("🚀 Start Matching", type="primary"):
        # Initialize progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Process resumes
        status_text.text("Processing resumes...")
        resume_data = []
        failed_files = []
        
        for i, file in enumerate(uploaded_files):
            status_text.text(f"Processing {file.name}...")
            
            resume_info, error_msg = process_resume(file)
            if resume_info:
                resume_data.append(resume_info)
            else:
                failed_files.append((file.name, error_msg))
            
            progress_bar.progress((i + 1) / len(uploaded_files))
        
        # Show processing results
        if failed_files:
            st.warning(f"⚠️ Failed to process {len(failed_files)} files:")
            for filename, error in failed_files:
                st.write(f"- {filename}: {error}")
        
        if not resume_data:
            st.error("❌ No valid resumes processed. Please check your files.")
            return
        
        if len(resume_data) < 2:
            st.info("ℹ️ Consider uploading more resumes for better comparison")
        
        st.success(f"✅ Successfully processed {len(resume_data)} resumes")
        
        # Build indexes
        status_text.text("Building search indexes...")
        progress_bar.progress(0.6)
        
        bm25_index = build_bm25_index(resume_data)
        if not bm25_index:
            st.error("❌ Failed to build BM25 index")
            return
        
        progress_bar.progress(0.8)
        
        vector_store = build_faiss_index(resume_data)
        if not vector_store:
            st.error("❌ Failed to build FAISS index")
            return
        
        # Calculate scores
        status_text.text("Calculating similarity scores...")
        progress_bar.progress(0.9)
        
        results = calculate_scores(resume_data, bm25_index, vector_store, job_description)
        
        progress_bar.progress(1.0)
        status_text.text("✅ Analysis complete!")
        
        if not results:
            st.error("❌ Failed to calculate scores")
            return
        
        # Display results
        st.markdown("---")
        display_results(results)
        
        # Create visualization
        st.markdown("---")
        st.subheader("📊 Visual Analysis")
        create_visualization(results)
        
        # Excel download
        st.markdown("---")
        create_excel_download(results)
        
        # Summary statistics
        st.markdown("---")
        st.subheader("📈 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Resumes", len(results))
        
        with col2:
            avg_score = sum(r['final_score'] for r in results) / len(results)
            st.metric("Average Score", f"{avg_score:.1f}%")
        
        with col3:
            best_score = max(r['final_score'] for r in results)
            st.metric("Best Match", f"{best_score}%")
        
        with col4:
            total_skills = sum(len(r['skills']) for r in results)
            st.metric("Total Skills Found", total_skills)

if __name__ == "__main__":
    main()